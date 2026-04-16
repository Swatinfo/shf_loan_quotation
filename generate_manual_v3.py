# -*- coding: utf-8 -*-
"""
Generate SHF Operational Manual v3.0
Replicates exact formatting from v2.0 + incorporates all improvements.
"""

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# ── Color constants ──
C_DARK     = '2D2D2D'
C_BLUE     = '1B3A6B'
C_ORANGE   = 'E8490F'
C_GRAY     = '666666'
C_WHITE    = 'FFFFFF'
C_DDDDD    = 'DDDDDD'
C_CCCCC    = 'CCCCCC'
C_LIGHT_ORANGE = 'FFF0EB'
C_LIGHT_BLUE   = 'E8F0FA'
C_LIGHT_GRAY   = 'F5F5F5'
C_LIGHT_RED    = 'FFEBEE'
C_LIGHT_GREEN  = 'E8F5E9'
C_LIGHT_YELLOW = 'FFFDE7'
C_RED_TEXT      = 'C62828'
C_GREEN_TEXT    = '1E7E34'
C_YELLOW_TEXT   = '8B6914'

# ── Font sizes (in Pt) ──
SZ_TITLE     = Pt(30)   # ~381000 emu
SZ_SUBTITLE  = Pt(13)
SZ_TAGLINE   = Pt(12)
SZ_SECTION   = Pt(14)   # ~177800 emu
SZ_TOC_TITLE = Pt(10)
SZ_TOC_DESC  = Pt(9)
SZ_BODY      = Pt(10)   # ~127000 emu / 133350 emu
SZ_BODY_SM   = Pt(10.5) # ~133350
SZ_KPI_MAIN  = Pt(10.5)
SZ_KPI_SUB   = Pt(9.5)  # ~120650
SZ_SUB_TITLE = Pt(9)    # ~114300

doc = Document()

# ── Page setup ──
section = doc.sections[0]
section.page_width  = Emu(7560310)
section.page_height = Emu(10692130)
section.top_margin    = Emu(1080135)
section.bottom_margin = Emu(914400)
section.left_margin   = Emu(914400)
section.right_margin  = Emu(914400)

# ── Style defaults ──
style = doc.styles['Normal']
style.font.name = 'Archivo'
style.font.size = Pt(10)
style.font.color.rgb = RGBColor.from_string(C_DARK)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)

# ══════════════════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ══════════════════════════════════════════════════════════════════

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for side, color in [('top', top), ('bottom', bottom), ('left', left), ('right', right), ('start', left), ('end', right)]:
        if color:
            b = parse_xml(f'<w:{side} {nsdecls("w")} w:val="single" w:sz="12" w:space="0" w:color="{color}"/>')
            borders.append(b)
    tcPr.append(borders)

def add_run(para, text, size=None, color=None, bold=None, font_name=None):
    """Add a formatted run to a paragraph."""
    run = para.add_run(text)
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.font.bold = bold
    if font_name:
        run.font.name = font_name
    return run

def add_para(text, size=SZ_BODY, color=C_DARK, bold=False, align=None, space_before=None, space_after=None):
    """Add a paragraph to the document."""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    add_run(p, text, size=size, color=color, bold=bold)
    return p

def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    # Remove existing borders
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)

def set_table_width(table, width_pct=100):
    """Set table to full width."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    existing = tblPr.find(qn('w:tblW'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblW)

def set_cell_width(cell, width_emu):
    """Set cell width."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{int(width_emu)}" w:type="dxa"/>')
    existing = tcPr.find(qn('w:tcW'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcW)

def set_cell_padding(cell, top=60, bottom=60, left=120, right=120):
    """Set cell padding."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(mar)

def make_section_header(num, title_en, title_gu=None):
    """Create a section header: orange number badge + blue title + gray subtitle."""
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    remove_table_borders(table)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, C_DARK)
    set_cell_padding(cell, top=120, bottom=120, left=200, right=200)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, f'{num:02d}  ', size=Pt(16), color=C_ORANGE, bold=True)
    add_run(p, title_en, size=SZ_SECTION, color=C_WHITE, bold=True)
    if title_gu:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p2, title_gu, size=SZ_SUB_TITLE, color=C_CCCCC, bold=False)

def make_sub_header(letter, title_en, title_gu=None):
    """Create a sub-header: A. / B. / C. with blue background."""
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    remove_table_borders(table)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, C_BLUE)
    set_cell_padding(cell, top=80, bottom=80, left=200, right=200)
    p = cell.paragraphs[0]
    add_run(p, f'{letter}.  {title_en}', size=SZ_BODY, color=C_WHITE, bold=True)
    if title_gu:
        p2 = cell.add_paragraph()
        add_run(p2, title_gu, size=SZ_SUB_TITLE, color=C_CCCCC, bold=False)

def make_bullet_list(items):
    """Create a bullet list with ▸ markers on light gray background."""
    table = doc.add_table(rows=len(items), cols=2)
    set_table_width(table)
    remove_table_borders(table)
    for i, item in enumerate(items):
        cell0 = table.rows[i].cells[0]
        cell1 = table.rows[i].cells[1]
        set_cell_shading(cell0, C_LIGHT_GRAY)
        set_cell_shading(cell1, C_LIGHT_GRAY)
        set_cell_width(cell0, 400)
        set_cell_padding(cell0, left=200)
        set_cell_padding(cell1, right=200)
        p0 = cell0.paragraphs[0]
        add_run(p0, '▸', size=SZ_BODY_SM, color=C_ORANGE)
        p1 = cell1.paragraphs[0]
        add_run(p1, item, size=SZ_BODY, color=C_DARK)

def make_numbered_steps(items):
    """Create numbered steps (1. 2. 3. ...)."""
    table = doc.add_table(rows=len(items), cols=2)
    set_table_width(table)
    remove_table_borders(table)
    for i, item in enumerate(items):
        cell0 = table.rows[i].cells[0]
        cell1 = table.rows[i].cells[1]
        set_cell_width(cell0, 400)
        set_cell_padding(cell0, left=200)
        set_cell_padding(cell1, right=200)
        p0 = cell0.paragraphs[0]
        add_run(p0, f'{i+1}.', size=SZ_BODY, color=C_ORANGE)
        p1 = cell1.paragraphs[0]
        add_run(p1, item, size=SZ_BODY, color=C_DARK)

def make_data_table(headers, rows, col_widths=None):
    """Create a data table with dark header row."""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    set_table_width(table)
    remove_table_borders(table)
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, C_DARK)
        set_cell_padding(cell, top=80, bottom=80, left=120, right=120)
        p = cell.paragraphs[0]
        add_run(p, h, size=SZ_BODY, color=C_WHITE, bold=True)
    # Data rows
    for ri, row in enumerate(rows):
        bg = C_WHITE if ri % 2 == 0 else C_LIGHT_GRAY
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            set_cell_shading(cell, bg)
            set_cell_padding(cell, top=60, bottom=60, left=120, right=120)
            p = cell.paragraphs[0]
            add_run(p, str(val), size=SZ_BODY, color=C_DARK)

def make_kpi_box(kpi_text, kpr_label, kpr_text):
    """Create a KPI/KPR info box."""
    table = doc.add_table(rows=1, cols=1) # changed to single column per original
    set_table_width(table)
    remove_table_borders(table)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, C_BLUE)
    set_cell_padding(cell, top=100, bottom=100, left=200, right=200)
    p = cell.paragraphs[0]
    add_run(p, 'KPI — મુખ્ય લક્ષ્યાંક', size=SZ_KPI_MAIN, color=C_WHITE, bold=True)
    p2 = cell.add_paragraph()
    add_run(p2, kpi_text, size=SZ_KPI_SUB, color=C_DDDDD)
    p3 = cell.add_paragraph()
    add_run(p3, f'{kpr_label}', size=SZ_KPI_MAIN, color=C_WHITE, bold=True)
    p4 = cell.add_paragraph()
    add_run(p4, kpr_text, size=SZ_KPI_SUB, color=C_DDDDD)

def make_notice_box(icon, text, bg_color, text_color, bold_first=False):
    """Create a notice/callout box (!, i, *, etc.)."""
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    remove_table_borders(table)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, bg_color)
    set_cell_padding(cell, top=100, bottom=100, left=200, right=200)
    p = cell.paragraphs[0]
    if bold_first:
        # Split at first newline or period for bold header
        parts = text.split('\n', 1) if '\n' in text else [text]
        add_run(p, f'{icon}  {parts[0]}', size=SZ_BODY, color=text_color, bold=True)
        if len(parts) > 1:
            p2 = cell.add_paragraph()
            add_run(p2, parts[1], size=SZ_BODY, color=text_color)
    else:
        add_run(p, f'{icon}  {text}', size=SZ_BODY, color=text_color)

def make_bilingual_notice(en_label, en_text, gu_label, gu_text, en_bg=C_LIGHT_BLUE, gu_bg=C_LIGHT_BLUE, en_color=C_BLUE, gu_color=C_BLUE):
    """Create paired EN/GU notice boxes."""
    # EN
    t1 = doc.add_table(rows=1, cols=1)
    set_table_width(t1)
    remove_table_borders(t1)
    c1 = t1.rows[0].cells[0]
    set_cell_shading(c1, en_bg)
    set_cell_padding(c1, top=80, bottom=80, left=200, right=200)
    p1 = c1.paragraphs[0]
    add_run(p1, f'[EN]  {en_label}', size=SZ_BODY, color=en_color, bold=True)
    p1b = c1.add_paragraph()
    add_run(p1b, en_text, size=SZ_BODY, color=en_color)
    # GU
    t2 = doc.add_table(rows=1, cols=1)
    set_table_width(t2)
    remove_table_borders(t2)
    c2 = t2.rows[0].cells[0]
    set_cell_shading(c2, gu_bg)
    set_cell_padding(c2, top=80, bottom=80, left=200, right=200)
    p2 = c2.paragraphs[0]
    add_run(p2, f'[GU]  {gu_label}', size=SZ_BODY, color=gu_color, bold=True)
    p2b = c2.add_paragraph()
    add_run(p2b, gu_text, size=SZ_BODY, color=gu_color)

def make_toc_row(num, title_en, desc):
    """Create a single TOC entry."""
    table = doc.add_table(rows=1, cols=3)
    set_table_width(table)
    remove_table_borders(table)
    c0 = table.rows[0].cells[0]
    c1 = table.rows[0].cells[1]
    c2 = table.rows[0].cells[2]
    set_cell_shading(c0, C_LIGHT_ORANGE)
    set_cell_width(c0, 500)
    set_cell_padding(c0, top=60, bottom=60, left=120, right=60)
    set_cell_padding(c1, top=60, bottom=60, left=60, right=60)
    set_cell_padding(c2, top=60, bottom=60, left=60, right=120)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p0, f'{num:02d}', size=SZ_BODY, color=C_ORANGE, bold=True)
    p1 = c1.paragraphs[0]
    add_run(p1, title_en, size=SZ_TOC_TITLE, color=C_BLUE, bold=True)
    p2 = c2.paragraphs[0]
    add_run(p2, desc, size=SZ_TOC_DESC, color=C_GRAY)

def make_doc_checklist(items):
    """Create a document checklist table."""
    headers = ['#', 'Document Name / દસ્તાવેજ નામ', 'Mandatory']
    rows = []
    for i, (name_en, name_gu, mandatory) in enumerate(items, 1):
        rows.append([str(i), f'{name_en} / {name_gu}', mandatory])
    make_data_table(headers, rows)

def page_break():
    """Insert a page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)

import docx.enum.text

# ══════════════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ══════════════════════════════════════════════════════════════════

# ── COVER PAGE ──
p1 = doc.add_paragraph()
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p1.paragraph_format.space_before = Pt(40)
add_run(p1, 'SHF ઓપરેશનલ માર્ગદર્શિકા', size=SZ_TITLE, color=C_DARK, bold=True)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p2, 'Shreenathji Home Finance — Business Operations Manual', size=SZ_SUBTITLE, color=C_GRAY)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(20)
add_run(p3, '"Shaping Happiness Forever"', size=SZ_TAGLINE, color=C_ORANGE)

# Cover info table
cover_table = doc.add_table(rows=2, cols=3)
set_table_width(cover_table)
remove_table_borders(cover_table)
cover_data = [
    [('માસિક ટાર્ગેટ', '₹35 કરોડ', C_LIGHT_ORANGE),
     ('બ્રાન્ચ', 'રાજકોટ / જામનગર', C_LIGHT_BLUE),
     ('સ્થાપક', 'ડેનીશ માલવિયા', C_LIGHT_GRAY)],
    [('Document Version', 'v3.0 — April 2026', C_LIGHT_BLUE),
     ('Classification', 'Internal & Confidential', C_LIGHT_RED),
     ('Effective Date', 'April 2026', C_LIGHT_GRAY)]
]
for ri, row in enumerate(cover_data):
    for ci, (label, value, bg) in enumerate(row):
        cell = cover_table.rows[ri].cells[ci]
        set_cell_shading(cell, bg)
        set_cell_padding(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        add_run(p, label, size=Pt(8), color=C_GRAY)
        p2 = cell.add_paragraph()
        add_run(p2, value, size=SZ_BODY, color=C_DARK, bold=True)

# ── VERSION CONTROL ──
vc_header = doc.add_table(rows=1, cols=1)
set_table_width(vc_header)
remove_table_borders(vc_header)
c = vc_header.rows[0].cells[0]
set_cell_shading(c, C_BLUE)
set_cell_padding(c, top=80, bottom=80, left=200, right=200)
p = c.paragraphs[0]
add_run(p, 'Version Control', size=SZ_BODY, color=C_WHITE, bold=True)
p2 = c.add_paragraph()
add_run(p2, 'દસ્તાવેજ સુધારણા ઇતિહાસ', size=SZ_SUB_TITLE, color=C_CCCCC)

make_data_table(
    ['Version', 'Date', 'Changes', 'Author'],
    [
        ['v1.0', 'Jan 2026', 'Initial document created', 'Denish Malviya'],
        ['v1.1', 'Mar 2026', 'HR Policy + Incentive Slab added', 'Denish Malviya'],
        ['v2.0', 'Apr 2026', 'CRM Guide, Loan Journey, Rejection Handling, Glossary added', 'Denish Malviya'],
        ['v3.0', 'Apr 2026', 'BDH + Office Employee roles, CRM Portal guide, CRM-aligned stages, Data Privacy, IT Policy, Complaint SOP, Gujarati doc names, Incentive clarification added', 'Denish Malviya'],
    ]
)

# ── TABLE OF CONTENTS ──
toc_header = doc.add_table(rows=1, cols=1)
set_table_width(toc_header)
remove_table_borders(toc_header)
c = toc_header.rows[0].cells[0]
set_cell_shading(c, C_DARK)
set_cell_padding(c, top=100, bottom=100, left=200, right=200)
p = c.paragraphs[0]
add_run(p, 'Table of Contents', size=SZ_SECTION, color=C_WHITE, bold=True)
p2 = c.add_paragraph()
add_run(p2, 'અનુક્રમ', size=SZ_SUB_TITLE, color=C_CCCCC)

# Part A — Roles
make_notice_box('📋', 'Part A — Roles & Responsibilities / ભૂમિકાઓ અને જવાબદારીઓ', C_LIGHT_BLUE, C_BLUE, bold_first=True)
toc_items = [
    (1, 'Junior Loan Advisor (Loaner – Fresher)', 'JD | SOP | Monthly Milestones | KPI | DVR Template'),
    (2, 'Bank Coordinator (Bank Dedicated Staff)', 'JD | SOP | KPI'),
    (3, 'CRM Executive (Customer Relationship)', 'JD | SOP | Call Script | KPI'),
    (4, 'Branch Manager (Rajkot / Jamnagar)', 'JD | SOP | KPI'),
    (5, 'Business Development Head (BDH)', 'JD | SOP | KPI'),
    (6, 'Office Employee (Operations)', 'JD | SOP | KPI'),
]
for num, title, desc in toc_items:
    make_toc_row(num, title, desc)

# Part B — Policies & Processes
make_notice_box('📋', 'Part B — Policies & Processes / નીતિઓ અને પ્રક્રિયાઓ', C_LIGHT_BLUE, C_BLUE, bold_first=True)
toc_items2 = [
    (7, 'Escalation Matrix', 'Level 1 → 2 → 3 with TAT definitions'),
    (8, 'HR Policy', 'Office Hours | Leave | Probation | Grade Structure'),
    (9, 'Loan Product Guide + Document Checklist', 'Salaried | Self-Employed | Partnership | PVT LTD'),
    (10, 'Loan File Journey / Process Flow', 'Inquiry → Document Collection → Login → Sanction → Disburse'),
    (11, 'SHF CRM Portal Guide', 'Daily Actions | Key Modules | CRM Rules'),
    (12, 'Rejection Handling Guide', '8 Reasons + Recovery Steps + Customer Script'),
    (13, 'Customer Complaint Handling SOP', 'Complaint Logging | Resolution | Escalation'),
]
for num, title, desc in toc_items2:
    make_toc_row(num, title, desc)

# Part C — Compensation & HR
make_notice_box('📋', 'Part C — Compensation & HR / વળતર અને HR', C_LIGHT_BLUE, C_BLUE, bold_first=True)
toc_items3 = [
    (14, 'Referral & DSA Commission Policy', 'DSA Tie-up | Payout Structure | Staff Referral'),
    (15, 'Incentive Structure', 'Disbursement Slab Table — Marginal, Uncapped'),
    (16, 'Office Rules & Discipline', '6 Core Rules + Digital Confidentiality Policy'),
    (17, 'Data Privacy & DPDPA Compliance', 'Customer Data | Staff Obligations | Device Policy'),
    (18, 'IT & Device Policy', 'Passwords | Lost Device | CRM Access'),
    (19, 'Appointment Letter Draft', 'Standard Template with Grade Structure'),
]
for num, title, desc in toc_items3:
    make_toc_row(num, title, desc)

# Part D — Reference
make_notice_box('📋', 'Part D — Quick Reference / ઝડપી સંદર્ભ', C_LIGHT_BLUE, C_BLUE, bold_first=True)
toc_items4 = [
    (20, 'Quick Reference Card', 'Print-Friendly Daily Summary'),
    (21, 'Glossary', '25+ Key Terms Explained for Freshers'),
]
for num, title, desc in toc_items4:
    make_toc_row(num, title, desc)

# ══════════════════════════════════════════════════════════════════
# SECTION 01 — Junior Loan Advisor
# ══════════════════════════════════════════════════════════════════
page_break()
make_section_header(1, 'Junior Loan Advisor  (Loaner – Fresher)', 'જુનિયર લોન એડવાઈઝર')

make_sub_header('A', 'Job Description', 'જોબ ડિસ્ક્રિપ્શન')
make_bullet_list([
    'માર્કેટમાં નવા ગ્રાહકો, બિલ્ડર્સ, પ્રોપર્ટી ડીલર્સ અને CA સાથે નેટવર્કિંગ કરવું.',
    'હોમ લોન, LAP, ઇન્ડસ્ટ્રીયલ અને પ્રોજેક્ટ લોન વિશે માર્કેટમાં જાગૃતિ લાવવી.',
    'ગ્રાહક પાસેથી ચેકલિસ્ટ મુજબ પ્રારંભિક દસ્તાવેજો એકત્રિત કરવા.',
    'DVR (Daily Visit Report) નિયમિત ભરીને ઓફિસ CRM માં અપડેટ કરવું.',
])

make_sub_header('B', 'SOP — Standard Operating Procedure', 'કાર્ય પ્રક્રિયા')
make_numbered_steps([
    'સવારનું પ્લાનિંગ: રોજ સવારે 10:00 વાગ્યે દિવસનું સાઈટ વિઝિટ લિસ્ટ તૈયાર કરવું.',
    'ફિલ્ડ વર્ક: રોજની ઓછામાં ઓછી 5 નવી મુલાકાતો લેવી.',
    'હેન્ડઓવર: લાવેલી ફાઈલો તે જ દિવસે સાંજે 5:00 થી 7:00 દરમિયાન ઓફિસમાં સબમિટ કરવી.',
    'DVR ભરવો: દરરોજ સાંજે CRM સોફ્ટવેરમાં વિઝિટ નોંધ, ઇન્ક્વાયરી સ્ટેટસ અને ફોલો-અપ ડેટ અપડેટ કરવા.',
])

make_sub_header('B1', 'Monthly Milestones', 'માસિક માઈલ્સ્ટોન')
make_data_table(
    ['Week', 'Target Activity', 'Minimum Output'],
    [
        ['Week 1', 'Market visits + new contacts', '15 new inquiries logged in CRM'],
        ['Week 2', 'File collection + document follow-up', '5 complete files submitted to office'],
        ['Week 3', 'Follow-up on pending files + new visits', '3 files logged in bank portal'],
        ['Week 4', 'Push for sanctions + generate new pipeline', '₹1.5 Cr+ login target achieved'],
    ]
)

make_sub_header('C', 'KPI & KPR  (Performance Metrics)')
make_kpi_box(
    'દર મહિને ₹5 થી ₹7 કરોડ login. Week-wise milestone ફોલો કરવા.',
    'KPR — રિવ્યુ માપદંડ',
    'દૈનિક DVR ભર્યો? CRM અપડેટ? ઇન્ક્વાયરી નોંધ? — ત્રણ KPR daily review.'
)

make_sub_header('D', 'DVR Template', 'ડેઈલી વિઝિટ રિપોર્ટ')
make_data_table(
    ['ક્રમ', 'ગ્રાહક / ડીલરનું નામ', 'સ્થળ / સાઈટ', 'ચર્ચાનો વિષય', 'ફોલો-અપ તારીખ'],
    [['1', '', '', '', ''], ['2', '', '', '', ''], ['3', '', '', '', ''], ['4', '', '', '', ''], ['5', '', '', '', '']]
)

# ══════════════════════════════════════════════════════════════════
# SECTION 02 — Bank Coordinator
# ══════════════════════════════════════════════════════════════════
page_break()
make_section_header(2, 'Bank Coordinator  (Bank Dedicated Staff – Fresher)', 'બેંક કોઓર્ડિનેટર')

make_sub_header('A', 'Job Description', 'જોબ ડિસ્ક્રિપ્શન')
make_bullet_list([
    'પોતાની ફાળવેલ બેંક (ICICI, Axis, HDFC/Kotak) ના પોર્ટલ પર ફાઈલો લોગિન કરવી.',
    'બેંકના ક્રેડિટ ઓફિસર, વ્હેલ્યુઅર અને વકીલ સાથે સતત સંકલન રાખવું.',
    'ફાઈલને સેન્ક્શન થી ડિસ્બર્સમેન્ટ સુધી ઝડપથી પહોંચાડવી.',
])

make_sub_header('B', 'SOP', 'કાર્ય પ્રક્રિયા')
make_numbered_steps([
    'બેંક વિઝિટ: રોજ સવારે 11:00 વાગ્યે બેંકમાં જઈને પેન્ડિંગ ફાઈલો પર ફોલો-અપ કરવો.',
    'ડેટા એન્ટ્રી: ફાઈલ મળ્યાના 12 કલાકમાં પોર્ટલ પર લોગિન અને રેફરન્સ નંબર જનરેટ કરવો.',
    'ફોલો-અપ: દરરોજ સાંજે બેંકની ક્વેરીઝનું લિસ્ટ CRM Executive ને આપવું.',
    'CRM અપડેટ: દૈનિક ધોરણે ફાઈલ સ્ટેટસ CRM સિસ્ટમમાં અપ-ટુ-ડેટ રાખવો.',
])

make_sub_header('C', 'KPI & KPR')
make_kpi_box(
    'ફાઈલ મળ્યાના 7–10 દિવસમાં સેન્ક્શન (TAT). 48 કલાકથી વધારે stuck → BM ને escalate.',
    'KPR — રિવ્યુ માપદંડ',
    'Login ભૂલ %  |  Sanction TAT  |  Bank relationship score (BM quarterly rating).'
)

# ══════════════════════════════════════════════════════════════════
# SECTION 03 — CRM Executive (renamed from CRM Madam)
# ══════════════════════════════════════════════════════════════════
page_break()
make_section_header(3, 'CRM Executive  (Customer Relationship)', 'કસ્ટમર રિલેશનશિપ એક્ઝિક્યુટિવ')

make_sub_header('A', 'Job Description', 'જોબ ડિસ્ક્રિપ્શન')
make_bullet_list([
    'ગ્રાહકોને ફાઈલના સ્ટેટસ વિશે ફોન દ્વારા નિયમિત જાણકારી આપવી.',
    'ગ્રાહકો પાસેથી પેન્ડિંગ દસ્તાવેજો મંગાવવા અને ગ્રાહક સંતોષ જાળવવો.',
    'બેંકની ક્વેરી 24 કલાકમાં ક્લોઝ કરવી.',
])

make_sub_header('B', 'SOP', 'કાર્ય પ્રક્રિયા')
make_numbered_steps([
    'અપડેટ્સ: દરરોજ તમામ પ્રોસેસ ફાઈલોના ગ્રાહકોને સ્ટેટસ કોલ્સ કરવા.',
    'ક્વેરી મેનેજમેન્ટ: બેંકની ક્વેરી 24 કલાકમાં ગ્રાહક પાસેથી સોલ્વ કરાવવી.',
    'પોસ્ટ-સેલ્સ: લોન મળ્યા બાદ ગ્રાહકને \'હેપીનેસ કોલ\' કરવો.',
])

make_sub_header('C', 'Call Script / Language Guide', 'કોલ સ્ક્રિપ્ટ')
make_data_table(
    ['પ્રસંગ', 'ભાષા / Script'],
    [
        ['ઓળખ', '"નમસ્કાર, હું SHF માંથી [નામ] બોલું છું. [ગ્રાહકનું નામ] સાથે વાત થઈ શકે?"'],
        ['સ્ટેટસ અપડેટ', '"તમારી ફાઈલ [સ્ટેટસ] સ્ટેજ પર છે. અંદાજિત [તારીખ] સુધી આગળ જશે."'],
        ['ડોક્યુ. માગ', '"બેંક તરફથી [X] ડોક્યુ. જોઈએ છે. WhatsApp પર મોકલી આપો."'],
        ['ક્વેરી delay', '"ગ્રાહકભાઈ, process માં સહેજ સમય લાગી શકે. અમે daily follow-up કરીએ છીએ."'],
        ['હેપીનેસ કોલ', '"અભિનંદન! તમારી લોન મંજૂર. SHF તરફથી ખૂબ ખૂબ શુભેચ્છાઓ!"'],
    ]
)

make_sub_header('D', 'KPI & KPR')
make_kpi_box(
    '100% ગ્રાહકોને સમયસર અપડેટ. 24 કલાકમાં ક્વેરી ક્લોઝ. Zero unanswered calls at day end.',
    'KPR — રિવ્યુ માપદંડ',
    'Customer satisfaction feedback  |  Query closure rate  |  Happiness call % of disbursed loans.'
)

# ══════════════════════════════════════════════════════════════════
# SECTION 04 — Branch Manager
# ══════════════════════════════════════════════════════════════════
page_break()
make_section_header(4, 'Branch Manager  (Rajkot / Jamnagar)', 'બ્રાન્ચ મેનેજર')

make_sub_header('A', 'Job Description', 'જોબ ડિસ્ક્રિપ્શન')
make_bullet_list([
    'આખી ટીમનું સુપરવિઝન અને ₹35 કરોડના ટાર્ગેટનું મોનિટરિંગ.',
    'ટીમમાં શિસ્ત જાળવવી અને નવા ફ્રેશર્સને માર્ગદર્શન આપવું.',
    'અઠવાડિક MIS રિપોર્ટ HO (Head Office) ને સબમિટ કરવો.',
])

make_sub_header('B', 'SOP', 'કાર્ય પ્રક્રિયા')
make_numbered_steps([
    'સવાર: 10:15 વાગ્યે \'મોર્નિંગ હડલ\' દ્વારા દિવસનું લક્ષ્ય નક્કી કરવું.',
    'સાંજ: 6:30 વાગ્યે MIS રિપોર્ટ (Login vs Disbursement) ચેક કરવો.',
    'શનિવાર: વીકલી રિવ્યુ અને ફ્રેશર્સ માટે સ્પેશ્યલ ટ્રેનિંગ સેશન.',
    'એસ્કેલેશન: કોઈ ફાઈલ 48 કલાક stuck → HO ને email + call.',
])

make_sub_header('C', 'KPI & KPR')
make_kpi_box(
    'માસિક ₹15–20 કરોડ disbursement. Staff attrition < 10%. Weekly MIS — every Monday 10 AM.',
    'KPR — રિવ્યુ માપદંડ',
    'Team KPI hit rate  |  Staff retention  |  Branch disbursement growth vs last month.'
)

# ══════════════════════════════════════════════════════════════════
# SECTION 05 — BDH (NEW)
# ══════════════════════════════════════════════════════════════════
page_break()
make_section_header(5, 'Business Development Head  (BDH)', 'બિઝનેસ ડેવલપમેન્ટ હેડ')

make_sub_header('A', 'Job Description', 'જોબ ડિસ્ક્રિપ્શન')
make_bullet_list([
    'બ્રાન્ચ લેવલના બિઝનેસ ડેવલપમેન્ટ ટાર્ગેટ પ્લાનિંગ અને મોનિટરિંગ.',
    'Loan Advisor ટીમના ફિલ્ડ વર્કનું સુપરવિઝન અને DVR રિવ્યુ.',
    'DSA / બિલ્ડર / CA ચેનલ પાર્ટનરશિપ ડેવલપ અને મેનેજ કરવી.',
    'બ્રાન્ચ ટીમના General Tasks અને Loan Tasks નું ઓવરસાઈટ (CRM દ્વારા).',
    'Branch Manager સાથે સંકલન કરીને માર્કેટ સ્ટ્રેટેજી ઘડવી.',
])

make_sub_header('B', 'SOP', 'કાર્ય પ્રક્રિયા')
make_numbered_steps([
    'સવાર: 10:00 વાગ્યે CRM Portal માં બ્રાન્ચ ડેશબોર્ડ ચેક — pending inquiries, overdue tasks, DVR status.',
    'ફિલ્ડ: Loan Advisors ની ફિલ્ડ વિઝિટ પ્લાન રિવ્યુ, ઓછામાં ઓછી 2 joint visits / week.',
    'સાંજ: DVR submissions ચેક, missing DVRs માટે advisor ને follow-up.',
    'Weekly: BM સાથે pipeline review — conversion rate, stuck files, advisor performance.',
    'Monthly: Channel partner (DSA/builder) relationship review, new tie-ups propose to HO.',
])

make_sub_header('C', 'KPI & KPR')
make_kpi_box(
    'બ્રાન્ચ monthly login target ₹10+ Cr. DSA channel contribution 20%+. Advisor DVR compliance 90%+.',
    'KPR — રિવ્યુ માપદંડ',
    'Pipeline conversion rate  |  Channel partner revenue %  |  Advisor activity compliance  |  Weekly BM report.'
)

# ══════════════════════════════════════════════════════════════════
# SECTION 06 — Office Employee (NEW)
# ══════════════════════════════════════════════════════════════════
page_break()
make_section_header(6, 'Office Employee  (Operations)', 'ઓફિસ કર્મચારી (ઓપરેશન્સ)')

make_sub_header('A', 'Job Description', 'જોબ ડિસ્ક્રિપ્શન')
make_bullet_list([
    'લોન ફાઈલોની ઓફિસ-સાઈડ પ્રોસેસિંગ: ટેક્નિકલ વેલ્યુએશન, ડોકેટ રિવ્યુ, OTC ક્લિયરન્સ.',
    'E-Sign અને eNACH ની ઓફિસ-સાઈડ ફેઝ હેન્ડલ કરવી (CRM workflow મુજબ).',
    'Sanction Letter અને Rate & PF ડોક્યુમેન્ટ્સ verify અને ફોરવર્ડ કરવા.',
    'ફાઈલ ડોક્યુમેન્ટ્સ ફિઝિકલ અને ડિજિટલ organize રાખવા.',
    'બેંક / Advisor / BM ને જરૂરી ઓફિસ સપોર્ટ પૂરો પાડવો.',
])

make_sub_header('B', 'SOP', 'કાર્ય પ્રક્રિયા')
make_numbered_steps([
    'સવાર: CRM Portal માં assigned tasks ચેક — valuation reports, docket reviews, OTC pending.',
    'પ્રોસેસિંગ: Assigned stage tasks same day complete — valuation report verify, docket documents check.',
    'ડોકેટ: Docket login phase — documents complete check, bank submission prepare.',
    'OTC: Cheque disbursement cases — OTC clearance documents verify, bank confirmation log.',
    'સાંજ: CRM status update — all assigned tasks progress, blockers BM ને report.',
])

make_sub_header('C', 'KPI & KPR')
make_kpi_box(
    'Assigned tasks same-day completion 95%+. Valuation/Docket TAT < 24 hours. Zero document errors.',
    'KPR — રિવ્યુ માપદંડ',
    'Task completion rate  |  Document accuracy  |  Stage TAT adherence  |  BM weekly review.'
)

# ══════════════════════════════════════════════════════════════════
# SECTION 07 — Escalation Matrix
# ══════════════════════════════════════════════════════════════════
page_break()
make_section_header(7, 'Escalation Matrix', 'એસ્કેલેશન મેટ્રિક્સ')
add_para('કોઈ પણ સ્ટાફ નીચેના ક્રમ પ્રમાણે ઉપર જઈ શકે — TAT definitions સ્પષ્ટ:', size=SZ_BODY, color=C_DARK)

make_data_table(
    ['Level', 'Trigger / Situation', 'TAT to Escalate', 'Contact', 'Response SLA'],
    [
        ['Level 1', 'Daily file/query stuck in process', '48 hours no update', 'Branch Manager', 'Same day'],
        ['Level 2', 'BM unresponsive / file stuck > 5 days / customer complaint', '24 hrs after L1', 'HO (Head Office)', 'Within 24 hours'],
        ['Level 3', 'Fraud suspicion / serious HR issue / legal', 'Immediate', 'Denish Malviya (Founder)', 'Immediate — call directly'],
    ]
)

make_notice_box('!', 'Rule: Never hide a problem. Escalation is NOT a punishment — it is a process. Staff who escalate on time are rewarded, not penalized.', C_LIGHT_YELLOW, C_YELLOW_TEXT, bold_first=True)

make_bilingual_notice(
    'English Notice',
    'If you are unsure whether to escalate — escalate. It is always better to inform your manager than to try to resolve alone and delay the process. Silence on a problem is a policy violation.',
    'ગુજરાતી નોંધ',
    'જો escalate કરવું કે નહીં તે અંગે confusion હોય — escalate કરો. Manager ને જાણ કરવી હંમેશા better છે. એકલા resolve કરવા attempt કરીને process delay કરવી policy violation ગણાય.'
)

# ══════════════════════════════════════════════════════════════════
# SECTION 08 — HR Policy
# ══════════════════════════════════════════════════════════════════
page_break()
make_section_header(8, 'HR Policy', 'HR નીતિ')

make_sub_header('A', 'Office Hours', 'ઓફિસ સમય')
make_bullet_list([
    'ઓફિસ સમય: સવારે 10:00 થી સાંજે 7:00',
    '10:15ની ટીમ મીટિંગ: ફરજિયાત હાજરી',
    '3 વખત મોડા આવ્યા = 1 દિવસ CL (Casual Leave) કાપ',
])

make_sub_header('B', 'Leave Policy', 'રજા નીતિ')
make_data_table(
    ['Leave Type', 'Per Year', 'Rule'],
    [
        ['Casual Leave (CL)', '12 days', 'Inform BM 1 day in advance'],
        ['Sick Leave (SL)', '6 days', '3+ days: Doctor certificate mandatory'],
        ['Earned Leave (EL)', '15 days', 'Apply 7 days in advance, BM approval required'],
        ['Emergency Leave', 'Up to 2 days', 'Inform BM within 2 hours — unpaid if not approved'],
    ]
)

make_sub_header('C', 'Probation & Grade Structure', 'પ્રોબેશન અને ગ્રેડ સ્ટ્રક્ચર')
make_data_table(
    ['Grade', 'Designation', 'Probation', 'Fixed Salary Band', 'Review Cycle'],
    [
        ['G1', 'Junior Loan Advisor (Fresher)', '3 months', '₹12,000 – ₹18,000', '3-month + Annual'],
        ['G2', 'Senior Loan Advisor (1+ yr)', 'Nil', '₹18,000 – ₹28,000', 'Annual'],
        ['G3', 'Bank Coordinator', '3 months', '₹15,000 – ₹22,000', '3-month + Annual'],
        ['G4', 'CRM Executive', '3 months', '₹14,000 – ₹20,000', '3-month + Annual'],
        ['G5', 'Branch Manager', 'Nil', '₹35,000 – ₹55,000', 'Annual'],
        ['G6', 'BDH (Business Development Head)', 'Nil', '₹30,000 – ₹45,000', 'Annual'],
        ['G7', 'Office Employee', '3 months', '₹14,000 – ₹22,000', '3-month + Annual'],
    ]
)

make_bullet_list([
    'KPI ન સંતોષ્યા: 1 મહિનો extension, ત્યારબાદ termination.',
    'Confirmed employee notice period: 1 month.',
    'Performance bonus: Annual — based on branch target achievement.',
])

# ══════════════════════════════════════════════════════════════════
# SECTION 09 — Loan Product Guide + Document Checklist
# ══════════════════════════════════════════════════════════════════
page_break()
make_section_header(9, 'Loan Product Guide + Document Checklist', 'લોન પ્રોડક્ટ ગાઈડ + ચેકલિસ્ટ')

make_sub_header('A', 'Loan Products at a Glance')
make_data_table(
    ['Product', 'Loan Amount', 'Tenure', 'Target Customer'],
    [
        ['Home Loan', '₹5L – ₹2Cr', '5–30 years', 'Salaried / Self-Employed buying residential property'],
        ['LAP (Loan Against Property)', '₹10L – ₹5Cr', '3–15 years', 'Business owners needing liquidity against owned property'],
        ['Industrial Loan', '₹25L – ₹10Cr', '5–20 years', 'Manufacturers / factories for plant & machinery'],
        ['Project Loan', '₹50L+', 'Case to Case', 'Builders / developers for project funding'],
    ]
)

make_notice_box('!', 'IMPORTANT NOTICE — Document Collection | દસ્તાવેજ સંગ્રહ સૂચના\nEnglish: All documents must be self-attested copies. Originals must be shown at bank visit. Incomplete files will NOT be accepted. Any missing document must be obtained within 48 hours.\nગુજરાતી: તમામ documents self-attested copy હોવી જોઈએ. Bank visit સમયે originals ફરજિયાત. અધૂરી ફાઈલ સ્વીકારાશે નહીં. Missing document 48 કલાકમાં મેળવવા.', C_LIGHT_RED, C_RED_TEXT, bold_first=True)

# B. Proprietor Checklist
make_sub_header('B', 'Document Checklist — Proprietor / Self-Employed', 'પ્રોપ્રાઈટર / સ્વ-રોજગાર')
make_doc_checklist([
    ('Passport Size Photographs Both', 'બંનેના પાસપોર્ટ સાઈઝ ફોટો', '✓ Yes'),
    ('PAN Card Both', 'બંનેના PAN કાર્ડ', '✓ Yes'),
    ('Aadhaar Card Both', 'બંનેના આધાર કાર્ડ', '✓ Yes'),
    ('GST Registration Certificate', 'GST રજિસ્ટ્રેશન સર્ટિફિકેટ', '✓ Yes'),
    ('Udyam Registration Certificate', 'ઉદ્યમ રજિસ્ટ્રેશન સર્ટિફિકેટ', '✓ Yes'),
    ('ITR (Last 3 years)', 'ITR (છેલ્લા 3 વર્ષ)', '✓ Yes'),
    ('Bank Statement (Last 12 months)', 'બેંક સ્ટેટમેન્ટ (છેલ્લા 12 મહિના)', '✓ Yes'),
    ('Current Loan Statement (if applicable)', 'ચાલુ લોન સ્ટેટમેન્ટ (લાગુ હોય તો)', 'If applicable'),
    ('Property File Xerox', 'પ્રોપર્ટી ફાઈલ ઝેરોક્ષ', '✓ Yes'),
])

# C. Partnership Checklist
make_sub_header('C', 'Document Checklist — Partnership / LLP', 'પાર્ટનરશિપ / LLP')
make_doc_checklist([
    ('PAN Card of Firm', 'ફર્મનું PAN કાર્ડ', '✓ Yes'),
    ('Partnership Deed', 'પાર્ટનરશિપ ડીડ', '✓ Yes'),
    ('GST Registration Certificate', 'GST રજિસ્ટ્રેશન સર્ટિફિકેટ', '✓ Yes'),
    ('ITR With Audit of Firm (Last 3 years)', 'ફર્મનું ઓડિટ સાથે ITR (છેલ્લા 3 વર્ષ)', '✓ Yes'),
    ('Firm Current A/c Bank Statement (Last 12 months)', 'ફર્મ ચાલુ ખાતા બેંક સ્ટેટમેન્ટ (છેલ્લા 12 મહિના)', '✓ Yes'),
    ('Current Loan Statement (if applicable)', 'ચાલુ લોન સ્ટેટમેન્ટ (લાગુ હોય તો)', 'If applicable'),
    ('Passport Size Photographs of All Partners', 'તમામ પાર્ટનર્સના પાસપોર્ટ સાઈઝ ફોટો', '✓ Yes'),
    ('PAN Card of All Partners', 'તમામ પાર્ટનર્સના PAN કાર્ડ', '✓ Yes'),
    ('Aadhaar Card of All Partners', 'તમામ પાર્ટનર્સના આધાર કાર્ડ', '✓ Yes'),
    ('ITR of Partners (Last 3 years)', 'પાર્ટનર્સના ITR (છેલ્લા 3 વર્ષ)', '✓ Yes'),
    ('Bank Statement of Partners (Last 12 months)', 'પાર્ટનર્સના બેંક સ્ટેટમેન્ટ (છેલ્લા 12 મહિના)', '✓ Yes'),
])

# D. PVT LTD Checklist
make_sub_header('D', 'Document Checklist — PVT LTD / Company', 'પ્રાઈવેટ લિમિટેડ / કંપની')
make_doc_checklist([
    ('PAN Card of Company', 'કંપનીનું PAN કાર્ડ', '✓ Yes'),
    ('Memorandum of Association (MOA)', 'મેમોરેન્ડમ ઓફ એસોસિએશન (MOA)', '✓ Yes'),
    ('Articles of Association (AOA)', 'આર્ટિકલ્સ ઓફ એસોસિએશન (AOA)', '✓ Yes'),
    ('GST Registration Certificate', 'GST રજિસ્ટ્રેશન સર્ટિફિકેટ', '✓ Yes'),
    ('ITR With Audit Report of Company (Last 3 years)', 'કંપનીનું ઓડિટ રિપોર્ટ સાથે ITR (છેલ્લા 3 વર્ષ)', '✓ Yes'),
    ('Current Loan Statement (if applicable)', 'ચાલુ લોન સ્ટેટમેન્ટ (લાગુ હોય તો)', 'If applicable'),
    ('Company Current A/c Statement (Last 12 months)', 'કંપની ચાલુ ખાતા સ્ટેટમેન્ટ (છેલ્લા 12 મહિના)', '✓ Yes'),
    ('Passport Size Photographs of All Directors', 'તમામ ડિરેક્ટર્સના પાસપોર્ટ સાઈઝ ફોટો', '✓ Yes'),
    ('PAN Card of All Directors', 'તમામ ડિરેક્ટર્સના PAN કાર્ડ', '✓ Yes'),
    ('Aadhaar Card of All Directors', 'તમામ ડિરેક્ટર્સના આધાર કાર્ડ', '✓ Yes'),
    ('ITR of Directors (Last 3 years)', 'ડિરેક્ટર્સના ITR (છેલ્લા 3 વર્ષ)', '✓ Yes'),
    ('Bank Statement of Directors (Last 12 months)', 'ડિરેક્ટર્સના બેંક સ્ટેટમેન્ટ (છેલ્લા 12 મહિના)', '✓ Yes'),
])

# E. Salaried Checklist
make_sub_header('E', 'Document Checklist — Salaried', 'પગારદાર')
make_doc_checklist([
    ('Passport Size Photographs Both', 'બંનેના પાસપોર્ટ સાઈઝ ફોટો', '✓ Yes'),
    ('PAN Card Both', 'બંનેના PAN કાર્ડ', '✓ Yes'),
    ('Aadhaar Card Both', 'બંનેના આધાર કાર્ડ', '✓ Yes'),
    ('Salary Slips (Last 6 months)', 'સેલેરી સ્લિપ (છેલ્લા 6 મહિના)', '✓ Yes'),
    ('ITR (Last 2 years)', 'ITR (છેલ્લા 2 વર્ષ)', '✓ Yes'),
    ('Form 16 (Last 2 years)', 'ફોર્મ 16 (છેલ્લા 2 વર્ષ)', '✓ Yes'),
    ('Bank Statement (Last 6 months)', 'બેંક સ્ટેટમેન્ટ (છેલ્લા 6 મહિના)', '✓ Yes'),
    ('Property Documents (if applicable)', 'પ્રોપર્ટી ડોક્યુમેન્ટ્સ (લાગુ હોય તો)', 'If applicable'),
])

make_notice_box('i', 'Note / નોંધ: The above document list matches the SHF CRM system (Settings → Documents tab). Any changes to the document list must be updated in the CRM by the Super Admin, and this manual must be revised accordingly. ઉપરની document list SHF CRM system સાથે match કરે છે. Document list માં ફેરફાર Super Admin દ્વારા CRM માં update કરવો અને manual revise કરવી.', C_LIGHT_BLUE, C_BLUE)

# ══════════════════════════════════════════════════════════════════
# SECTION 10 — Loan File Journey (UPDATED to match CRM stages)
# ══════════════════════════════════════════════════════════════════
page_break()
make_section_header(10, 'Loan File Journey — Process Flow', 'લોન ફાઈલ પ્રોસેસ ફ્લો')
add_para('દરેક ફાઈલ નીચેના stages માંથી પસાર થાય. SHF CRM Portal માં દરેક stage track થાય છે:', size=SZ_BODY, color=C_DARK)

make_sub_header('A', 'Simplified Process (Staff View)')
make_data_table(
    ['Stage #', 'Stage Name', 'Who is Responsible', 'TAT', 'Output'],
    [
        ['1', 'Lead Generation / Inquiry', 'Junior Loan Advisor', 'Ongoing', 'New inquiry in CRM'],
        ['2', 'Document Collection', 'Junior Loan Advisor', '2–3 days from lead', 'Complete file folder'],
        ['3', 'File Submission to Office', 'Junior Loan Advisor', 'Same day as collection', 'File handed to BM'],
        ['4', 'File Review & Login Prep', 'Branch Manager', 'Within 24 hours', 'File approved for login'],
        ['5', 'Bank Portal Login', 'Bank Coordinator', 'Within 12 hrs of BM approval', 'Login reference number'],
        ['6', 'Bank Query Management', 'Bank Coordinator + CRM Executive', 'Within 24 hrs of query', 'Query resolved'],
        ['7', 'Valuation / Legal', 'Bank Coordinator + Office Employee', '5–7 days', 'Valuation + Legal report received'],
        ['8', 'Sanction Letter', 'Bank Coordinator', 'Day 7–10 from login', 'Sanction letter to customer'],
        ['9', 'Customer Acceptance', 'CRM Executive', 'Within 48 hrs of sanction', 'Signed sanction copy'],
        ['10', 'Disbursement', 'Bank Coordinator + Office Employee', 'As per bank schedule', 'Amount credited'],
        ['11', 'Happiness Call', 'CRM Executive', 'Same day as disbursement', 'Customer satisfaction logged'],
    ]
)

make_sub_header('B', 'Detailed CRM Stages (System View)')
make_notice_box('i', 'The SHF CRM Portal tracks each file through more granular stages. Staff should be familiar with these CRM stages:', C_LIGHT_BLUE, C_BLUE, bold_first=True)
make_data_table(
    ['CRM Stage', 'Description', 'Key Role(s)'],
    [
        ['1. Inquiry', 'New lead entry, customer details capture', 'Loan Advisor'],
        ['2. Document Selection', 'Select required documents based on customer type', 'Loan Advisor'],
        ['3. Document Collection', 'Collect & verify all required documents', 'Loan Advisor'],
        ['4a. Application Number', 'Bank portal login, get reference number', 'Bank Coordinator'],
        ['4b. BSM/OSV', 'Bank site manager / on-site verification', 'Bank Coordinator'],
        ['4c. Legal Verification', '3-phase: Advisor → Bank → Office Employee', 'Multi-role handoff'],
        ['4d. Technical Valuation', 'Property valuation by authorized valuer', 'Office Employee'],
        ['4e. Sanction Decision', 'Bank approves / escalates / rejects', 'Bank Coordinator'],
        ['5. Rate & PF', '3-phase: Rate finalization, processing fee', 'Multi-role handoff'],
        ['6. Sanction Letter', '3-phase: Generate, verify, deliver', 'Multi-role handoff'],
        ['7. Docket Login', '3-phase: Document compile, bank submit', 'Multi-role handoff'],
        ['8. KFS', 'Key Facts Statement to customer', 'Bank Coordinator'],
        ['9. E-Sign & eNACH', '4-phase: Digital sign & auto-debit setup', 'Multi-role handoff'],
        ['10. Disbursement', 'Final amount transfer', 'Bank Coordinator'],
        ['11. OTC Clearance', 'Cheque disbursement cases only', 'Office Employee'],
    ]
)

make_notice_box('!', 'TAT Alert: If any stage exceeds its TAT by 48+ hours — mandatory escalation to Branch Manager. BM logs it in the weekly MIS report.', C_LIGHT_YELLOW, C_YELLOW_TEXT, bold_first=True)

# ══════════════════════════════════════════════════════════════════
# SECTION 11 — CRM Software Guide (UPDATED — filled placeholder)
# ══════════════════════════════════════════════════════════════════
page_break()
make_section_header(11, 'SHF CRM Portal Guide', 'SHF CRM પોર્ટલ ગાઈડ')

make_notice_box('i', 'SHF uses a custom-built CRM Portal (SHF CRM). All staff must be trained in Week 1 of joining. Login credentials are issued by the Branch Manager on Day 1. Access the portal at the URL provided by your Branch Manager.', C_LIGHT_BLUE, C_BLUE)

make_sub_header('A', 'Daily Mandatory Actions in CRM')
make_data_table(
    ['Time', 'Action', 'Who', 'Where in CRM'],
    [
        ['10:00 AM', 'Log today\'s planned site visits', 'Junior Advisor', 'DVR Section → New Visit'],
        ['During field', 'Update each visit outcome in real-time', 'Junior Advisor', 'DVR → Edit Entry'],
        ['12:00 PM', 'Log any new bank queries received', 'Bank Coordinator', 'Loan → Stage → Query'],
        ['5:00–7:00 PM', 'Submit files / update file status', 'All Staff', 'Loan Dashboard → Status Update'],
        ['7:00 PM', 'Mark daily DVR as complete', 'Junior Advisor', 'DVR → Submit'],
        ['7:00 PM', 'Send CRM query list to CRM Executive', 'Bank Coordinator', 'Queries → Export / WhatsApp'],
    ]
)

make_sub_header('B', 'Key CRM Modules')
make_data_table(
    ['Module', 'Purpose', 'Access Level'],
    [
        ['Dashboard', 'Overview of active loans, tasks, pending queries, stats', 'All Staff'],
        ['Quotations', 'Create bank comparison PDFs for customers (EMI, charges)', 'Loan Advisors, BM, BDH'],
        ['Loans', 'Track each file from inquiry to disbursement', 'All Staff'],
        ['DVR (Daily Visit Report)', 'Log daily field visits, follow-ups, outcomes', 'Loan Advisors, BDH, BM'],
        ['General Tasks', 'Personal/delegated tasks — reminders, follow-ups', 'All Staff'],
        ['Notifications', 'Bell icon alerts for stage changes, queries, assignments', 'All Staff'],
        ['Settings', 'Banks, documents, stages, user management', 'Admin / Super Admin only'],
    ]
)

make_sub_header('C', 'CRM Rules')
make_bullet_list([
    'CRM is the single source of truth — verbal updates are NOT valid. Everything must be logged.',
    'No screenshots of CRM to be shared outside official WhatsApp groups.',
    'CRM login is personal — never share your password with anyone.',
    'If CRM is down: maintain a manual Excel log and sync within 2 hours of system restoration.',
])

make_bilingual_notice(
    'CRM Policy Notice',
    'Every customer interaction, file update, and bank query must be logged in CRM on the same day. End-of-day CRM audit is done by Branch Manager. Staff with incomplete CRM entries will receive a warning. Third warning = written notice.',
    'CRM નીતિ નોંધ',
    'દરેક customer interaction, file update, અને bank query — same day CRM માં log કરવી ફરજિયાત. Branch Manager દ્વારા day-end CRM audit. Incomplete entries = warning. ત્રીજી warning = written notice.'
)

# ══════════════════════════════════════════════════════════════════
# SECTION 12 — Rejection Handling Guide
# ══════════════════════════════════════════════════════════════════
page_break()
make_section_header(12, 'Rejection Handling Guide', 'રિજેક્શન હેન્ડલિંગ ગાઈડ')
add_para('Bank rejection is common. Here is how to handle each type professionally:', size=SZ_BODY, color=C_DARK)

make_sub_header('A', 'Common Rejection Reasons & Recovery Steps')
make_data_table(
    ['Rejection Reason', 'Why It Happens', 'Recovery Action', 'Who Acts'],
    [
        ['Low CIBIL Score (< 650)', 'Past EMI defaults / high credit utilization', 'Apply after 6 months clean payment. Suggest credit improvement.', 'CRM Executive + Customer'],
        ['Income Insufficient', 'EMI-to-income ratio > 50%', 'Add co-applicant. Reduce loan amount.', 'Advisor + BM'],
        ['Property Title Issue', 'Disputed ownership / missing documents', 'Engage lawyer. Get clear title certificate.', 'Bank Coord + BM'],
        ['Business vintage < 2 years', 'Self-employed business too new', 'Explore NBFC or lower amount. Re-apply after 2 years.', 'BM + Advisor'],
        ['Multiple ongoing loans', 'High FOIR', 'Close smaller loans first. Apply with reduced amount.', 'BM + Customer'],
        ['Property valuation low', 'Market value less than expected', 'Appeal with 2nd valuer. Top-up with own contribution.', 'Bank Coordinator'],
        ['Incomplete documents', 'Missing or incorrect documents', 'Re-collect correct documents within 48 hours. Resubmit.', 'Junior Advisor'],
        ['Negative area/property', 'Property in bank\'s negative list', 'Switch bank. Try NBFC/HFC.', 'Bank Coord + BM'],
    ]
)

make_sub_header('B', 'What To Tell the Customer')
# DO say
t1 = doc.add_table(rows=1, cols=1)
set_table_width(t1)
remove_table_borders(t1)
c1 = t1.rows[0].cells[0]
set_cell_shading(c1, C_LIGHT_GREEN)
set_cell_padding(c1, top=80, bottom=80, left=200, right=200)
p1 = c1.paragraphs[0]
add_run(p1, '[EN]  DO Say (English)', size=SZ_BODY, color=C_GREEN_TEXT, bold=True)
p1b = c1.add_paragraph()
add_run(p1b, '"We understand this is disappointing. We are reviewing the exact reason and will suggest the best path forward within 24 hours."', size=SZ_BODY, color=C_GREEN_TEXT)

t1g = doc.add_table(rows=1, cols=1)
set_table_width(t1g)
remove_table_borders(t1g)
c1g = t1g.rows[0].cells[0]
set_cell_shading(c1g, C_LIGHT_GREEN)
set_cell_padding(c1g, top=80, bottom=80, left=200, right=200)
p1g = c1g.paragraphs[0]
add_run(p1g, '[GU]  આ કહો (ગુજરાતી)', size=SZ_BODY, color=C_GREEN_TEXT, bold=True)
p1gb = c1g.add_paragraph()
add_run(p1gb, '"અમને સમજાય છે કે આ disappointing છે. અમે exact reason review કરી રહ્યા છીએ અને 24 કલાકમાં શ્રેષ્ઠ રસ્તો suggest કરીશું."', size=SZ_BODY, color=C_GREEN_TEXT)

# NEVER say
t2 = doc.add_table(rows=1, cols=1)
set_table_width(t2)
remove_table_borders(t2)
c2 = t2.rows[0].cells[0]
set_cell_shading(c2, C_LIGHT_RED)
set_cell_padding(c2, top=80, bottom=80, left=200, right=200)
p2 = c2.paragraphs[0]
add_run(p2, '[EN]  NEVER Say (English)', size=SZ_BODY, color=C_RED_TEXT, bold=True)
p2b = c2.add_paragraph()
add_run(p2b, '"The bank has rejected you" — Never say this without context. Never blame the customer. Never promise approval on the next attempt without first confirming with BM.', size=SZ_BODY, color=C_RED_TEXT)

t2g = doc.add_table(rows=1, cols=1)
set_table_width(t2g)
remove_table_borders(t2g)
c2g = t2g.rows[0].cells[0]
set_cell_shading(c2g, C_LIGHT_RED)
set_cell_padding(c2g, top=80, bottom=80, left=200, right=200)
p2g = c2g.paragraphs[0]
add_run(p2g, '[GU]  ક્યારેય ન કહો (ગુજરાતી)', size=SZ_BODY, color=C_RED_TEXT, bold=True)
p2gb = c2g.add_paragraph()
add_run(p2gb, '"બેંકે તમને reject કર્યા" — આ ક્યારેય context વિના ન કહો. ગ્રાહકને ક્યારેય blame ન કરો. BM સાથે confirm કર્યા વિના next attempt ઉપર approval promise ન કરો.', size=SZ_BODY, color=C_RED_TEXT)

make_sub_header('C', 'Rejection Escalation Rules')
make_bullet_list([
    'All rejections must be logged in CRM within 4 hours of receiving rejection letter.',
    'If rejection reason unclear: Bank Coordinator calls Credit Officer same day for clarification.',
    'BM reviews all rejections weekly — identifies patterns and alerts HO.',
    'If 3+ rejections from same bank in same week: BM calls bank relationship manager immediately.',
])

make_bilingual_notice(
    'Customer Handling Notice',
    'A rejected customer, handled well, often becomes a successful customer through a different bank or after improving their profile. Never lose the relationship. Follow up within 30 days to check if their situation has improved.',
    'ગ્રાહક સંભાળ નોંધ',
    'Rejection મળેલ ગ્રાહક, સારી રીતે handle કરો, different bank અથવા profile improve કરીને successful customer બની શકે. Relationship ક્યારેય lose ન કરો. 30 days માં follow-up કરો.',
    en_bg=C_LIGHT_GREEN, gu_bg=C_LIGHT_GREEN, en_color=C_GREEN_TEXT, gu_color=C_GREEN_TEXT
)

# ══════════════════════════════════════════════════════════════════
# SECTION 13 — Customer Complaint Handling SOP (NEW)
# ══════════════════════════════════════════════════════════════════
page_break()
make_section_header(13, 'Customer Complaint Handling SOP', 'ગ્રાહક ફરિયાદ હેન્ડલિંગ SOP')

make_sub_header('A', 'Complaint Logging', 'ફરિયાદ નોંધણી')
make_numbered_steps([
    'ગ્રાહકની ફરિયાદ CRM Portal માં "General Task" તરીકે log કરો — category: "Customer Complaint".',
    'ફરિયાદ receive થયાના 1 કલાકમાં Branch Manager ને notify કરો.',
    'ફરિયાદના type identify કરો: Service Delay / Document Issue / Mis-communication / Charges Dispute / Other.',
    'ગ્રાહકને acknowledgment call/message — "તમારી ફરિયાદ register થઈ છે. 24 કલાકમાં resolution."',
])

make_sub_header('B', 'Resolution Matrix', 'નિરાકરણ માળખું')
make_data_table(
    ['Complaint Type', 'Resolution Owner', 'TAT', 'Escalation If Unresolved'],
    [
        ['Service Delay', 'Bank Coordinator', '24 hours', 'BM → HO (48 hrs)'],
        ['Document Issue', 'Junior Advisor', '24 hours', 'BM (24 hrs)'],
        ['Mis-communication', 'CRM Executive', '12 hours', 'BM (24 hrs)'],
        ['Charges Dispute', 'BM + Bank Coordinator', '48 hours', 'HO → Founder (72 hrs)'],
        ['Staff Behavior', 'Branch Manager', '24 hours', 'HO → Founder (48 hrs)'],
    ]
)

make_sub_header('C', 'Post-Resolution', 'નિરાકરણ પછી')
make_numbered_steps([
    'Resolution complete થયા બાદ CRM task close with resolution notes.',
    'CRM Executive દ્વારા ગ્રાહકને satisfaction call — "શું તમારી ફરિયાદ સંતોષકારક રીતે ઉકેલાઈ?"',
    'Monthly complaint report BM → HO: total complaints, resolved %, patterns identified.',
])

make_bilingual_notice(
    'Complaint Policy',
    'Every complaint is an opportunity to strengthen the customer relationship. No complaint should go unacknowledged beyond 1 hour. Zero tolerance for staff who ignore or hide customer complaints.',
    'ફરિયાદ નીતિ',
    'દરેક ફરિયાદ ગ્રાહક સંબંધ મજબૂત કરવાની તક છે. 1 કલાકથી વધુ unacknowledged complaint ન રહેવી જોઈએ. ગ્રાહક ફરિયાદ ignore કરનાર સ્ટાફ સામે zero tolerance.'
)

# ══════════════════════════════════════════════════════════════════
# SECTION 14 — Referral & DSA Commission Policy
# ══════════════════════════════════════════════════════════════════
page_break()
make_section_header(14, 'Referral & DSA Commission Policy', 'રેફરલ અને DSA કમિશન નીતિ')

make_sub_header('A', 'Who is a DSA?')
add_para('DSA (Direct Selling Agent) = External individual or firm who refers loan customers to SHF. They are registered partners, not SHF employees.', size=SZ_BODY, color=C_DARK)

make_sub_header('B', 'DSA Tie-Up Process', 'DSA ટાઈ-અપ પ્રક્રિયા')
make_numbered_steps([
    'Interested DSA meets Branch Manager and submits KYC + business proof.',
    'BM reviews and forwards to HO for approval.',
    'HO approves and assigns DSA Code within 5 working days.',
    'DSA Agreement signed — commission structure shared in writing.',
    'DSA starts referring: all files must come with DSA Code on the file cover.',
])

make_sub_header('C', 'DSA Commission Payout Structure', 'DSA કમિશન સ્ટ્રક્ચર')
make_data_table(
    ['Loan Type', 'Commission %', 'Paid On', 'Payout Timeline'],
    [
        ['Home Loan', '0.30% – 0.50%', 'Disbursed amount', 'Within 15 days of disbursement'],
        ['LAP', '0.40% – 0.60%', 'Disbursed amount', 'Within 15 days of disbursement'],
        ['Industrial / Project', '0.25% – 0.40%', 'Disbursed amount', 'Within 20 days of disbursement'],
    ]
)

make_notice_box('i', 'Commission paid only after successful disbursement — not on sanction alone. TDS as per Income Tax Act will be deducted. DSA must provide PAN for payment.', C_LIGHT_BLUE, C_BLUE)

make_sub_header('D', 'Internal Staff Referral Policy', 'સ્ટાફ રેફરલ')
make_bullet_list([
    'Staff can refer external customers — referral must be declared to BM.',
    'Internal referral bonus: ₹500 per disbursed file (paid with monthly salary).',
    'Staff cannot act as DSA for any other finance company while employed at SHF.',
])

# ══════════════════════════════════════════════════════════════════
# SECTION 15 — Incentive Structure (UPDATED — clarified as marginal)
# ══════════════════════════════════════════════════════════════════
page_break()
make_section_header(15, 'Incentive Structure  (Slab Table)', 'ઇન્સેન્ટિવ સ્ટ્રક્ચર')
add_para('માસિક ડિસ્બર્સમેન્ટ પ્રમાણે ઇન્સેન્ટિવ (ઇન્ડિવિઝ્યુઅલ સ્ટાફ):', size=SZ_BODY, color=C_DARK)

make_notice_box('i', 'Calculation Method: MARGINAL SLAB (like Income Tax)\nEach slab percentage applies ONLY to the amount within that slab, not the entire disbursement.\nExample: If you disburse ₹6 Cr → First ₹3Cr at 0.10% (₹30,000) + Next ₹2Cr at 0.15% (₹30,000) + Next ₹1Cr at 0.20% (₹20,000) = Total ₹80,000.', C_LIGHT_BLUE, C_BLUE, bold_first=True)

make_data_table(
    ['Disbursement Slab', 'Incentive %', 'Estimated Incentive', 'Cumulative Example'],
    [
        ['₹0 – ₹3 Cr', '0.10%', '₹0 – ₹30,000', '—'],
        ['₹3 – ₹5 Cr', '0.15%', '₹30,000 – ₹75,000', '₹3Cr done → ₹30K + next slab'],
        ['₹5 – ₹7 Cr', '0.20%', '₹75,000 – ₹1,40,000', '₹5Cr done → ₹60K + next slab'],
        ['₹7 Cr+', '0.25%', '₹1,40,000+', 'No ceiling — uncapped'],
    ]
)

make_bullet_list([
    'Slab calculated on individual staff\'s disbursed files only.',
    'Incentive paid on 5th of following month along with salary.',
    'If staff leaves before month-end, incentive pro-rated based on disbursed amount till last working day.',
    'Branch Manager incentive calculated separately on total branch disbursement.',
])

# ══════════════════════════════════════════════════════════════════
# SECTION 16 — Office Rules & Discipline
# ══════════════════════════════════════════════════════════════════
page_break()
make_section_header(16, 'Office Rules & Discipline', 'ઓફિસ શિસ્ત અને નિયમો')

make_data_table(
    ['Rule', 'Description'],
    [
        ['(1)  સમયપાલન', 'સવારે 10:00 થી સાંજે 7:00. 10:15ની ટીમ મીટિંગ ફરજિયાત. 3 late = 1 CL deduction.'],
        ['(2)  ડ્રેસ કોડ', 'ઓફિસ અને field visit — formal dress mandatory. No casuals, no slippers.'],
        ['(3)  પારદર્શિતા', 'ગ્રાહકને હંમેશા સાચી અને સ્પષ્ટ માહિતી. No false promises on loan amounts or timelines.'],
        ['(4)  ઝીરો પેન્ડન્સી', 'Today\'s file = today\'s CRM update. No overnight pending without BM approval.'],
        ['(5)  ટીમ વર્ક', 'ક્વેરી આવે → BM ને તત્કાળ જાણ. Hiding problems = misconduct.'],
        ['(6)  ગોપનીયતા (Digital)', 'PROHIBITED: Screenshots of customer files | Forwarding docs on personal WhatsApp | Sharing CRM data outside official groups. Violation = immediate termination.'],
    ]
)

make_bilingual_notice(
    'Misconduct Warning',
    'Any staff found sharing customer documents outside official channels, violating confidentiality, making false promises to customers, or hiding problems from management will face immediate disciplinary action up to and including termination without notice.',
    'ગેરવર્તણૂક ચેતવણી',
    'Customer documents share કરવા, confidentiality નો ભંગ, ગ્રાહકને ખોટા વચન, management થી problem છૂપાવવી — આ ગેરવર્તણૂક ગણાય. Immediate disciplinary action — Termination without notice સુધી.',
    en_bg=C_LIGHT_RED, gu_bg=C_LIGHT_RED, en_color=C_RED_TEXT, gu_color=C_RED_TEXT
)

# ══════════════════════════════════════════════════════════════════
# SECTION 17 — Data Privacy & DPDPA Compliance (NEW)
# ══════════════════════════════════════════════════════════════════
page_break()
make_section_header(17, 'Data Privacy & DPDPA Compliance', 'ડેટા ગોપનીયતા અને DPDPA અનુપાલન')

make_notice_box('!', 'India\'s Digital Personal Data Protection Act (DPDPA), 2023 applies to SHF. All staff handling customer data must comply.', C_LIGHT_RED, C_RED_TEXT, bold_first=True)

make_sub_header('A', 'What is Personal Data?', 'વ્યક્તિગત ડેટા શું છે?')
make_bullet_list([
    'Customer name, address, phone number, email — all personal data.',
    'PAN, Aadhaar, bank statements, ITR — sensitive personal data.',
    'Property documents, loan details, CIBIL score — financial personal data.',
    'Any information that can identify a specific person = personal data.',
])

make_sub_header('B', 'Staff Obligations', 'સ્ટાફ ફરજો')
make_numbered_steps([
    'ગ્રાહકના દસ્તાવેજો ફક્ત loan processing માટે use — personal purpose માટે ક્યારેય નહીં.',
    'ગ્રાહકના ડેટાની ક copies personal device (phone/laptop) માં save ન કરવી.',
    'WhatsApp/email દ્વારા documents share — ફક્ત official groups/email IDs દ્વારા.',
    'ગ્રાહકે consent આપ્યો હોય તે purpose માટે જ data use કરવો.',
    'Data breach (leaked documents, unauthorized access) — BM ને immediately report.',
])

make_sub_header('C', 'Penalties for Violation', 'ઉલ્લંઘન માટે દંડ')
make_data_table(
    ['Violation', 'Internal Action', 'Legal Risk'],
    [
        ['Sharing documents on personal WhatsApp', 'Written warning → Termination', 'DPDPA fine up to ₹250 Cr on company'],
        ['Saving customer data on personal device', 'Written warning → Data wipe', 'Company liability under DPDPA'],
        ['Using customer data for personal benefit', 'Immediate termination', 'Criminal prosecution possible'],
        ['Not reporting a data breach', 'Termination + internal inquiry', 'DPDPA mandates breach reporting'],
    ]
)

# ══════════════════════════════════════════════════════════════════
# SECTION 18 — IT & Device Policy (NEW)
# ══════════════════════════════════════════════════════════════════
page_break()
make_section_header(18, 'IT & Device Policy', 'IT અને ડિવાઈસ નીતિ')

make_sub_header('A', 'Password & Access Rules', 'પાસવર્ડ અને એક્સેસ નિયમો')
make_bullet_list([
    'CRM login password minimum 8 characters — mix of letters, numbers, and symbols.',
    'Password ક્યારેય colleague સાથે share ન કરો. Not even with BM.',
    'Password દર 90 days માં change — CRM system reminder આપશે.',
    'Public computer / cyber cafe થી CRM access ન કરવો.',
    'Browser માં "Remember Password" use ન કરો — especially shared devices.',
])

make_sub_header('B', 'Lost / Stolen Device', 'ખોવાયેલ / ચોરાયેલ ડિવાઈસ')
make_numbered_steps([
    'BM ને immediately (1 કલાકમાં) call/WhatsApp કરીને જાણ કરો.',
    'BM IT Admin (HO) ને inform → CRM session terminate + password reset.',
    'FIR / police complaint file (if phone with customer data).',
    'New device setup — fresh CRM login with new password.',
])

make_sub_header('C', 'Approved Apps & Communication', 'માન્ય Apps અને સંદેશાવ્યવહાર')
make_data_table(
    ['Purpose', 'Approved Tool', 'NOT Allowed'],
    [
        ['Customer communication', 'Official WhatsApp Business / Call', 'Personal WhatsApp for doc sharing'],
        ['File sharing (internal)', 'CRM Portal upload', 'Google Drive / Telegram / email attachments'],
        ['Status updates', 'CRM Portal', 'Verbal-only updates'],
        ['Bank queries', 'CRM Portal → Query module', 'Direct calls without CRM logging'],
    ]
)

# ══════════════════════════════════════════════════════════════════
# SECTION 19 — Appointment Letter Draft
# ══════════════════════════════════════════════════════════════════
page_break()
make_section_header(19, 'Appointment Letter — Draft Template', 'એપોઈન્ટમેન્ટ લેટર')

# Date/Place
t = doc.add_table(rows=1, cols=2)
set_table_width(t)
remove_table_borders(t)
c1 = t.rows[0].cells[0]
c2 = t.rows[0].cells[1]
p1 = c1.paragraphs[0]
add_run(p1, 'તારીખ:', size=SZ_BODY, color=C_DARK, bold=True)
add_run(p1, '  _______________', size=SZ_BODY, color=C_GRAY)
p2 = c2.paragraphs[0]
add_run(p2, 'સ્થળ:', size=SZ_BODY, color=C_DARK, bold=True)
add_run(p2, '  _______________', size=SZ_BODY, color=C_GRAY)
# Date/Place English
p_date = c1.add_paragraph()
add_run(p_date, 'Date:', size=SZ_BODY, color=C_GRAY)
add_run(p_date, '  _______________', size=SZ_BODY, color=C_GRAY)
p_place = c2.add_paragraph()
add_run(p_place, 'Place:', size=SZ_BODY, color=C_GRAY)
add_run(p_place, '  _______________', size=SZ_BODY, color=C_GRAY)

add_para('', space_after=Pt(4))
add_para('પ્રતિ, / To,', size=SZ_BODY, color=C_DARK)
add_para('[ઉમેદવારનું નામ / Candidate Name]', size=SZ_BODY, color=C_DARK)
add_para('[સરનામું / Address]', size=SZ_BODY, color=C_DARK)
add_para('', space_after=Pt(4))
add_para('વિષય / Subject: [ડેઝિગ્નેશન / Designation] — નિમણૂક પત્ર / Appointment Letter', size=SZ_BODY, color=C_DARK, bold=True)
add_para('', space_after=Pt(4))
add_para('પ્રિય [ઉમેદવારનું નામ] / Dear [Candidate Name],', size=SZ_BODY, color=C_DARK)
add_para('અમને આનંદ છે કે SHF (Shreenathji Home Finance) માં [ડેઝિગ્નેશન] તરીકે તમારી નિમણૂક થઈ છે. અમારું વિઝન "Shaping Happiness Forever" છે.', size=SZ_BODY, color=C_DARK)
add_para('We are pleased to appoint you as [Designation] at SHF (Shreenathji Home Finance). Our vision is "Shaping Happiness Forever".', size=SZ_BODY, color=C_GRAY)
add_para('', space_after=Pt(4))
add_para('નિમણૂકની શરતો / Terms of Appointment:', size=SZ_BODY, color=C_DARK, bold=True)

make_data_table(
    ['Particulars', 'Details'],
    [
        ['Designation', '[e.g. Junior Loan Advisor]'],
        ['Grade', '[e.g. G1 / G2 / G3 / G4 / G5 / G6 / G7]'],
        ['Department', 'Home Loan / Operations'],
        ['Branch', 'Rajkot / Jamnagar'],
        ['Date of Joining', '_______________'],
        ['Fixed Monthly Salary', '₹_______________'],
        ['Incentive', 'As per Incentive Slab (Section 15 of this Manual)'],
        ['Probation Period', '3 months (G1, G3, G4, G7)  /  Nil (G2, G5, G6)'],
        ['Notice Period (Post-Confirmation)', '1 Month'],
        ['Office Hours', '10:00 AM – 7:00 PM, Monday to Saturday'],
        ['Reporting To', '[Branch Manager Name]'],
    ]
)

add_para('', space_after=Pt(8))
add_para('અમને વિશ્વાસ છે કે તમારા જોડાવાથી SHF ની પ્રગતિમાં નવો વેગ આવશે.', size=SZ_BODY, color=C_DARK)
add_para('We are confident that your joining will bring new momentum to SHF\'s growth.', size=SZ_BODY, color=C_GRAY)
add_para('', space_after=Pt(4))
add_para('શુભેચ્છાઓ સહ, / Best wishes,', size=SZ_BODY, color=C_DARK)
add_para('', space_after=Pt(8))
add_para('_______________________________', size=SZ_BODY, color=C_DARK)
add_para('(ડેનીશ માલવિયા / Denish Malviya)', size=SZ_BODY, color=C_DARK, bold=True)
add_para('સ્થાપક, SHF — Shreenathji Home Finance / Founder, SHF', size=SZ_BODY, color=C_GRAY)

# ══════════════════════════════════════════════════════════════════
# SECTION 20 — Quick Reference Card
# ══════════════════════════════════════════════════════════════════
page_break()
make_section_header(20, 'Quick Reference Card  (Print & Keep)', 'ઝડપી સંદર્ભ કાર્ડ')

make_notice_box('*', 'This page is designed to be printed and kept at your desk. All key information at a glance.', C_LIGHT_YELLOW, C_YELLOW_TEXT)

# Daily Timings
make_notice_box('►', 'Daily Timings', C_LIGHT_BLUE, C_BLUE, bold_first=True)
make_data_table(
    ['Time', 'Activity'],
    [
        ['10:00 AM', 'Reach office. Log today\'s DVR plan in CRM.'],
        ['10:15 AM', 'Morning Huddle — Mandatory. No exceptions.'],
        ['11:00 AM', 'Bank Coordinator: Visit assigned bank for file follow-ups.'],
        ['5:00–7:00 PM', 'Submit collected files to office. Update CRM status.'],
        ['7:00 PM', 'Close day: DVR submitted, all queries updated, BM informed of blockers.'],
    ]
)

# Monthly Targets
make_notice_box('►', 'Monthly Targets at a Glance', C_LIGHT_BLUE, C_BLUE, bold_first=True)
make_data_table(
    ['Role', 'Monthly Target', 'Weekly Milestone'],
    [
        ['Junior Loan Advisor', '₹5–7 Cr login', 'Week 1: 15 leads  |  Week 4: ₹1.5Cr+ login'],
        ['Bank Coordinator', 'TAT < 10 days', 'All queries closed within 24 hours'],
        ['CRM Executive', '100% customer updates', 'Zero pending queries at week end'],
        ['Branch Manager', '₹15–20 Cr disbursement', 'Weekly MIS every Monday 10 AM to HO'],
        ['BDH', '₹10+ Cr branch login', 'DSA channel 20%+ contribution'],
        ['Office Employee', '95%+ same-day task completion', 'Zero document errors per week'],
    ]
)

# Key Contacts
make_notice_box('►', 'Key Contacts', C_LIGHT_BLUE, C_BLUE, bold_first=True)
make_data_table(
    ['Person', 'Role', 'Contact'],
    [
        ['Denish Malviya', 'Founder / Level 3 Escalation', '+91 99747 89089'],
        ['[BM Name — fill per branch]', 'Branch Manager — Rajkot', '[Phone Number]'],
        ['[BM Name — fill per branch]', 'Branch Manager — Jamnagar', '[Phone Number]'],
        ['Head Office', 'Level 2 Escalation', '[HO Email / Phone]'],
    ]
)
make_notice_box('i', 'Each branch must fill in the BM name and contact number above before distributing this card. / દરેક branch ને ઉપર BM નામ અને contact number ભરવા.', C_LIGHT_BLUE, C_BLUE)

# Escalation Quick Guide
make_notice_box('►', 'Escalation Quick Guide', C_LIGHT_BLUE, C_BLUE, bold_first=True)
make_data_table(
    ['If this happens...', 'Do this...'],
    [
        ['File stuck > 48 hours', '→ Escalate to Branch Manager immediately'],
        ['Bank rejection received', '→ Log in CRM within 4 hours. Inform BM same day.'],
        ['Customer complaint', '→ Log in CRM as General Task. Inform BM within 1 hour. Do not argue with customer.'],
        ['CRM down', '→ Manual Excel log. Sync within 2 hours of restoration.'],
        ['Fraud / suspicious activity', '→ Call Denish Malviya directly. Do not wait.'],
        ['Lost phone with CRM access', '→ Call BM immediately. Password reset within 1 hour.'],
    ]
)

# ══════════════════════════════════════════════════════════════════
# SECTION 21 — Glossary (EXPANDED)
# ══════════════════════════════════════════════════════════════════
page_break()
make_section_header(21, 'Glossary — Key Terms Explained', 'ગ્લોસેરી — મુખ્ય શબ્દો')
add_para('For new staff (Freshers): Read this section first before joining field work.', size=SZ_BODY, color=C_DARK)

make_data_table(
    ['Term', 'Full Form', 'Meaning (English)', 'અર્થ (ગુજરાતી)'],
    [
        ['KYC', 'Know Your Customer', 'Customer identity verification — PAN, Aadhaar, photo.', 'ગ્રાહકની ઓળખ ચકાસણી — PAN, Aadhaar, ફોટો.'],
        ['DVR', 'Daily Visit Report', 'Daily log of all field visits made by Loan Advisor.', 'Loan Advisor દ્વારા કરેલ field visits ની દૈનિક નોંધ.'],
        ['CRM', 'Customer Relationship Management', 'Software used to track all customers, files, and queries.', 'Software જે ગ્રાહકો, ફાઈલો અને queries track કરે.'],
        ['LAP', 'Loan Against Property', 'Loan given by keeping your owned property as collateral.', 'પોતાની property ગીરો મૂકીને મળતી loan.'],
        ['TAT', 'Turnaround Time', 'Maximum time allowed to complete a task or stage.', 'કોઈ કામ અથવા stage પૂર્ણ કરવા માટે મળેલ maximum સમય.'],
        ['MIS', 'Management Information System', 'Weekly report of login vs disbursement data sent to HO.', 'Login vs Disbursement ડેટાનો weekly report — HO ને મોકલાય.'],
        ['DSA', 'Direct Selling Agent', 'External partner who refers loan customers to SHF.', 'External partner જે SHF ને loan ગ્રાહકો refer કરે.'],
        ['CIBIL', 'Credit Information Bureau India Ltd', 'Credit score bureau — score 300–900. Lenders check this.', 'Credit score bureau — score 300–900. Bank loan approve કરતા પહેલા check કરે.'],
        ['EMI', 'Equated Monthly Instalment', 'Fixed monthly loan repayment amount.', 'દર મહિને ભરવામાં આવતી fixed loan repayment રકમ.'],
        ['FOIR', 'Fixed Obligation to Income Ratio', '% of income going toward existing EMIs. Banks prefer < 50%.', 'Income નો % જે existing EMIs ભરવામાં જાય. Bank 50% થી ઓછો prefer કરે.'],
        ['Sanction', '—', 'Bank officially approves the loan — sanction letter issued.', 'Bank officially loan approve કરે — sanction letter મળે.'],
        ['Disbursement', '—', 'Bank actually transfers the loan amount to customer / seller.', 'Bank actually loan ની રકમ ગ્રાહક / seller ને transfer કરે.'],
        ['Login', '—', 'Submitting a complete file to the bank portal for processing.', 'Complete file bank portal ઉપર process માટે submit કરવી.'],
        ['BM', 'Branch Manager', 'In-charge of entire branch operations and team.', 'Branch ની તમામ operations અને team ના ઇન-ચાર્જ.'],
        ['BDH', 'Business Development Head', 'Senior role overseeing branch business targets and advisor team.', 'Branch business targets અને advisor team ના ઓવરસાઈટ માટેની senior ભૂમિકા.'],
        ['HO', 'Head Office', 'Main SHF office — escalation Level 2.', 'SHF ની main office — escalation Level 2 નો સંપર્ક.'],
        ['ITR', 'Income Tax Return', 'Annual income declaration filed with Income Tax Department.', 'Income Tax Department ને ભરવામાં આવતો વાર્ષિક income declaration.'],
        ['GST', 'Goods and Services Tax', 'Business tax registration certificate.', 'Business ના tax registration નું certificate.'],
        ['LTV', 'Loan to Value Ratio', '% of property value that bank will fund. Typically 75–80%.', 'Property value નો % જે bank fund કરશે. સામાન્ય રીતે 75–80%.'],
        ['NBFC', 'Non-Banking Financial Company', 'Financial institutions (not banks) that also give loans. e.g. Bajaj Finance.', 'Financial institutions (bank નહીં) જે loan આપે. દા.ત. Bajaj Finance.'],
        ['HFC', 'Housing Finance Company', 'Specialized NBFC for home loans. e.g. LIC Housing.', 'Home loan માટે specialized NBFC. દા.ત. LIC Housing.'],
        ['DPDPA', 'Digital Personal Data Protection Act', 'India\'s 2023 data privacy law — applies to all businesses handling personal data.', 'ભારતનો 2023 ડેટા privacy કાયદો — personal data handle કરતા તમામ businesses ને લાગુ.'],
        ['OTC', 'Over The Counter', 'Physical cheque clearance for loan disbursement (if not digital).', 'Loan disbursement માટે physical cheque clearance (digital ન હોય ત્યારે).'],
        ['eNACH', 'Electronic National Automated Clearing House', 'Auto-debit mandate setup for EMI collection from customer bank.', 'Customer bank માંથી EMI collection માટે auto-debit mandate setup.'],
        ['KFS', 'Key Facts Statement', 'RBI-mandated document showing all loan terms clearly to customer before signing.', 'RBI-mandated document — signing પહેલા ગ્રાહકને loan ની તમામ terms સ્પષ્ટ દર્શાવે.'],
    ]
)

# ── SAVE ──
output_path = 'SHF_Operational_Manual_2026_v3.docx'
doc.save(output_path)
print(f'Document saved: {output_path}')
print('Done!')
